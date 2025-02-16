from docx import Document
import openpyxl, os
import comtypes.client  # Para convertir Word a PDF en Windows
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns

class GenerarCertificado:
    # 📌 Inicializar la clase con el DataFrame de accionistas
    def __init__(self):
        #self.df_accionista = df_accionista
        self.path_ = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
        print(self.path_)
        self.path_documentos = os.path.join(self.path_, "Documentos")
        self.path_plantillas = os.path.join(self.path_documentos, "Plantillas")
        self.path_plantilla_excel = None
        self.path_plantilla_word = os.path.join(self.path_plantillas, "plantilla.docx")

        self.path_procesados = os.path.join(self.path_, "Procesados")
        self.path_certificados = os.path.join(self.path_procesados, "Certificados")
        self.path_pdf = None
        self.path_word = None
    # 📌 Generar certificado con los datos de un accionista
    def generarCertificado(self, nombre_cerificado):
        self.path_pdf = os.path.join(self.path_certificados,"Certificados_pdf", f"{nombre_cerificado}.pdf")
        self.path_word = os.path.join(self.path_certificados, "Certificados_word", f"{nombre_cerificado}.docx")

        self.__manejo_word(nombre_cerificado)
        self.__word_a_pdf(self.path_word, self.path_pdf)

        print(f"✅ PDF guardado como: {nombre_cerificado}.pdf")
        #return f"Certifico que {self.nombre} {self.apellido} ha realizado el curso de {self.curso} con fecha {self.fecha}"
    # 📌 Convertir Word a PDF
    def __word_a_pdf(self,input_path, output_path):
        word = comtypes.client.CreateObject("Word.Application")
        word.Visible = False
        doc = word.Documents.Open(input_path)
        # ❌ No se modifican los márgenes para respetar la plantilla con membrete
        doc.SaveAs(output_path, FileFormat=17)  # 17 es PDF en Word
        doc.Close()
        word.Quit()
    # 📌 Función para formatear números correctamente
    def __formato_numero(self, value, format_type):
        if format_type == "money":
            return f"${value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
            #return f"{value:.2f}".replace(",", "X").replace(".", ",").replace("X", ".")  # 📌 Formato con separación de miles con punto y decimal con coma
            #return f"${value:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")
        elif format_type == "percentage":
            return f"{value:.2f}%".replace(".", ",")
        elif format_type == "id":
            return f"{int(value):,}".replace(",", ".")  # Solo número con puntos, sin $
        return str(value)
    # 📌 Función para manejar el documento de Word
    def __manejo_word_der(self):
        # 📌 Cargar datos de Excel
        wb = openpyxl.load_workbook(self.path_plantilla_excel, data_only=True)
        ws = wb.active
        
        # 📌 Cargar la plantilla de Word
        doc = Document(self.path_plantilla_word)

        # 📌 Verificar si hay tablas en el documento
        if len(doc.tables) == 0:
            print("❌ Error: La plantilla no tiene tablas. Agrega una tabla en Word y vuelve a intentarlo.")
        else:
            table = doc.tables[0]  # Usar la primera tabla

        # 📌 Agregar datos a la tabla
        for row_idx, row in enumerate(ws.iter_rows()):
            values = [cell.value if cell.value is not None else "" for cell in row]
            # 📌 Unir las primeras 4 filas en ambas columnas con negrita y centradas
            if row_idx < 4:
                row_cells = table.add_row().cells
                cell = row_cells[0]
                cell.merge(row_cells[1])  # Fusionar ambas columnas
                cell.text = " ".join(str(cell.value) if cell.value is not None else "" for cell in row)

                # 📌 Formato centrado y negrita
                para = cell.paragraphs[0]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = para.runs[0]
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "Century Gothic"

            # 📌 La fila 5 ocupa ambas columnas pero ahora está alineada a la izquierda
            elif row_idx == 5:
                row_cells = table.add_row().cells
                cell = row_cells[0]
                cell.merge(row_cells[1])  # Fusionar ambas columnas
                cell.text = " ".join(str(cell.value) if cell.value is not None else "" for cell in row)

                # 📌 Formato alineado a la izquierda sin negrita
                para = cell.paragraphs[0]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                run = para.runs[0]
                run.font.size = Pt(9)
                run.font.name = "Century Gothic"
            elif row_idx == 26:

                row_cells = table.add_row().cells
                for i, value in enumerate(values):
                    row_cells[i].text = str(value)
                    para = row_cells[i].paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    run = para.runs[0]
                    run.font.size = Pt(8)
                    run.font.name = "Calibri"
                    run.bold = True
            
            elif 30 <= row_idx <= 32:
                # 📌 Hacer negrita las filas 31, 32 y 33
                row_cells = table.add_row().cells
                for i, value in enumerate(values):
                    row_cells[i].text = str(value)
                    para = row_cells[i].paragraphs[0]
                    run = para.runs[0]
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
                    run.bold = True  # Aplicar negrita

            elif 6 <= row_idx <= 12:
                # 📌 Alinear la columna A a la derecha y la columna B a la izquierda
                row_cells = table.add_row().cells
                for i, cell in enumerate(row):
                    value = cell.value
                    if isinstance(value, (int, float)):
                        if (row_idx == 7 and i == 1) or (row_idx == 9 and i == 1):  # 📌 Celdas B7 y B9
                            value = self.__formato_numero(value, "id")  # Solo número con puntos, sin $
                        elif "%" in cell.number_format:
                            value = self.__formato_numero(value * 100, "percentage")  # Convertir a porcentaje
                        else:
                            value = self.__formato_numero(value, "money")  # Formato de dinero
                    else:
                        value = str(value) if value is not None else ""
                    
                    row_cells[i].text = value
                    para = row_cells[i].paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT if i == 0 else WD_PARAGRAPH_ALIGNMENT.LEFT
                    run = para.runs[0]
                    run.font.size = Pt(9)
                    run.font.name = "Century Gothic"
                    if i == 1:
                        run.bold = True
            
            else:
                row_cells = table.add_row().cells
                for i, cell in enumerate(row):
                    value = cell.value
                    if isinstance(value, (int, float)):
                        if (row_idx == 7 and i == 1) or (row_idx == 9 and i == 1):  # 📌 Celdas B7 y B9
                            value = self.__formato_numero(value, "id")  # Solo número con puntos, sin $
                        elif "%" in cell.number_format:
                            value = self.__formato_numero(value * 100, "percentage")  # Convertir a porcentaje
                        else:
                            value = self.__formato_numero(value, "money")  # Formato de dinero
                    else:
                        value = str(value) if value is not None else ""

                    row_cells[i].text = value
                    para = row_cells[i].paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    run = para.runs[0]
                    run.font.size = Pt(9)
                    run.font.name = "Century Gothic"

                    if i == 1 and row_idx >= 6:
                        run.bold = True
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

        # 📌 Reducir espaciado entre párrafos para optimizar espacio
        for para in doc.paragraphs:
            pPr = para._element.get_or_add_pPr()
            spacing = OxmlElement("w:spacing")
            spacing.set(ns.qn("w:after"), "25")
            pPr.append(spacing)

        # 📌 Guardar el documento modificado
        doc.save(self.path_word)

    def __manejo_word_izq_der(self):
        # 📌 Cargar datos de Excel
        wb = openpyxl.load_workbook(self.path_plantilla_excel, data_only=True)
        ws = wb.active
        # 📌 Cargar la plantilla de Word
        doc = Document(self.path_plantilla_word)
        # 📌 Verificar si hay tablas en el documento
        if len(doc.tables) == 0:
            print("❌ Error: La plantilla no tiene tablas. Agrega una tabla en Word y vuelve a intentarlo.")
        else:
            table = doc.tables[0]  # Usar la primera tabla
        # 📌 Agregar datos a la tabla
        for row_idx, row in enumerate(ws.iter_rows()):
            values = [cell.value if cell.value is not None else "" for cell in row]
            # 📌 Unir las primeras 4 filas en ambas columnas con negrita y centradas
            if row_idx < 4:
                row_cells = table.add_row().cells
                cell = row_cells[0]
                cell.merge(row_cells[1])  # Fusionar ambas columnas
                cell.text = " ".join(str(cell.value) if cell.value is not None else "" for cell in row)
                # 📌 Formato centrado y negrita
                para = cell.paragraphs[0]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = para.runs[0]
                run.bold = True
                run.font.size = Pt(9)  # 📌 Reducido a tamaño 9
                run.font.name = "Century Gothic"
            # 📌 La fila 5 ocupa ambas columnas pero ahora está alineada y justificada
            elif row_idx == 5:
                row_cells = table.add_row().cells
                cell = row_cells[0]
                cell.merge(row_cells[1])  # Fusionar ambas columnas
                cell.text = " ".join(str(cell.value) if cell.value is not None else "" for cell in row)
                # 📌 Formato alineado a la izquierda sin negrita
                para = cell.paragraphs[0]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                run = para.runs[0]
                run.font.size = Pt(9)
                run.font.name = "Century Gothic"
            # 📌 
            elif row_idx == 26:
                row_cells = table.add_row().cells
                for i, value in enumerate(values):
                    row_cells[i].text = str(value)
                    para = row_cells[i].paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    run = para.runs[0]
                    run.font.size = Pt(8)
                    run.font.name = "Calibri"
                    run.bold = True
            # 📌 Filas 31, 32 y 33 en negrita
            elif 30 <= row_idx <= 32:
                row_cells = table.add_row().cells
                for i, value in enumerate(values):
                    row_cells[i].text = str(value)
                    para = row_cells[i].paragraphs[0]
                    run = para.runs[0]
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
                    run.bold = True  # Aplicar negrita
            # 📌 resto
            else:
                row_cells = table.add_row().cells
                for i, cell in enumerate(row):
                    value = cell.value
                    if isinstance(value, (int, float)):
                        # 📌 Aplicar formato correcto
                        # 📌 Celdas B7 y B9
                        if (row_idx == 7 and i == 1) or (row_idx == 9 and i == 1): 
                            # Solo número con puntos, sin $
                            value = self.__formato_numero(value, "id")  
                        elif "%" in cell.number_format:
                            # Convertir a porcentaje
                            value = self.__formato_numero(value * 100, "percentage")  
                        else:
                             # Formato de dinero
                            value = self.__formato_numero(value, "money") 
                    else:
                        value = str(value) if value is not None else ""
                    # 📌 Asignar texto y ajustar formato
                    row_cells[i].text = value
                    para = row_cells[i].paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    run = para.runs[0]
                    run.font.size = Pt(9)  # 📌 Reducido a tamaño 9
                    run.font.name = "Century Gothic"
                    # 📌 Poner en negrita los valores de la columna B desde la fila 7
                    if i == 1 and row_idx >= 6:
                        run.bold = True
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # 📌 Reducir espaciado entre párrafos para optimizar espacio
        for para in doc.paragraphs:
            pPr = para._element.get_or_add_pPr()
            spacing = OxmlElement("w:spacing")
            spacing.set(ns.qn("w:after"), "25")  # 📌 Espaciado reducido
            pPr.append(spacing)
        # 📌 Guardar el documento modificado
        doc.save(self.path_word)

    def __manejo_word(self, nombre_cerificado):
        self.path_plantilla_excel = os.path.join(self.path_certificados, "Certificados_excel",f"{nombre_cerificado}.xlsx")
        # 📌 Cargar datos de Excel
        wb = openpyxl.load_workbook(self.path_plantilla_excel, data_only=True)
        ws = wb.active
        # 📌 Cargar la plantilla de Word
        doc = Document(self.path_plantilla_word)
        # 📌 Verificar si hay tablas en el documento
        if len(doc.tables) == 0:
            print("❌ Error: La plantilla no tiene tablas. Agrega una tabla en Word y vuelve a intentarlo.")
            return
        else:
            table = doc.tables[0]  # Usar la primera tabla existente sin agregar filas adicionales
        # 📌 Agregar datos a la tabla
        for row_idx, row in enumerate(ws.iter_rows()):
            values = [cell.value if cell.value is not None else "" for cell in row]
            # 📌 Unir las primeras 4 filas en ambas columnas con negrita y centradas
            if row_idx < 4:
                # 📌 Unir las primeras 6 filas en ambas columnas con negrita y centradas
                row_cells = table.add_row().cells
                cell = row_cells[0]
                cell.merge(row_cells[1])
                cell.text = " ".join(map(str, values))
                para = cell.paragraphs[0]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                run = para.runs[0]
                run.bold = True
                run.font.size = Pt(9)
                run.font.name = "Century Gothic"
            # 📌 La fila 5 ocupa ambas columnas pero ahora está alineada a la justificada
            elif row_idx == 5:
                # 📌 Fusionar columnas A y B en la fila 6
                row_cells = table.add_row().cells
                cell = row_cells[0]
                cell.merge(row_cells[1])
                cell.text = " ".join(map(str, values))
                para = cell.paragraphs[0]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                run = para.runs[0]
                run.font.size = Pt(9)
                run.font.name = "Century Gothic"
            # 📌 Alinear la columna A a la derecha y la columna B a la izquierda
            elif 6 <= row_idx <= 12:
                # 📌 Fusionar columnas A y B manteniendo formatos correctos
                row_cells = table.add_row().cells
                value_a = str(values[0])
                value_b = values[1]

                if isinstance(value_b, (int, float)):
                    if row_idx in [7, 9]:
                        # 📌 Formato con separación de miles con punto y decimal con coma
                        value_b = f"{value_b:,}".replace(",", "X").replace(".", ",").replace("X", ".")
                    elif "%" in row[1].number_format:
                        # 📌 Formato de porcentaje
                        value_b = f"{value_b * 100:.2f}%"
                    else:
                        # 📌 Formato de moneda
                        value_b = self.__formato_numero(value_b, "money") 

                cell = row_cells[0]
                cell.merge(row_cells[1])
                para = cell.paragraphs[0]
                para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

                run_a = para.add_run(value_a + " ")
                run_a.font.size = Pt(9)
                run_a.font.name = "Century Gothic"

                run_b = para.add_run(str(value_b))
                run_b.bold = True
                run_b.font.size = Pt(9)
                run_b.font.name = "Century Gothic"
            # 📌 Filas 25 en Calibri 8
            elif row_idx == 26:
                # 📌 Fila 24 en Calibri 8
                row_cells = table.add_row().cells
                for i, value in enumerate(values):
                    row_cells[i].text = str(value)
                    para = row_cells[i].paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    run = para.runs[0]
                    run.font.size = Pt(8)
                    run.font.name = "Calibri"
                    run.bold = True
            # 📌 Filas 31, 32 y 33 en negrita
            elif 30 <= row_idx <= 32:
                # 📌 Hacer negrita las filas 31, 32 y 33
                row_cells = table.add_row().cells
                for i, value in enumerate(values):
                    row_cells[i].text = str(value)
                    para = row_cells[i].paragraphs[0]
                    run = para.runs[0]
                    run.font.size = Pt(9)
                    run.font.name = "Calibri"
                    run.bold = True  # Aplicar negrita
            # 📌 Filas restantes
            else:
                # 📌 Agregar filas sin duplicarlas ni agregar adicionales
                row_cells = table.add_row().cells
                for i, value in enumerate(values):
                    if isinstance(value, (int, float)):
                        if "%" in row[i].number_format:
                            value = f"{value * 100:.2f}%"
                        else:
                            value = self.__formato_numero(value, "money")
                    row_cells[i].text = str(value)
                    para = row_cells[i].paragraphs[0]
                    para.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    run = para.runs[0]
                    run.font.size = Pt(9)
                    run.font.name = "Century Gothic"
                    if i == 1 and row_idx >= 6:
                        run.bold = True
                        para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        # 📌 Reducir espaciado entre párrafos
        for para in doc.paragraphs:
            pPr = para._element.get_or_add_pPr()
            spacing = OxmlElement("w:spacing")
            spacing.set(ns.qn("w:after"), "25")
            pPr.append(spacing)
        # 📌 Guardar el documento modificado
        doc.save(self.path_word)

   